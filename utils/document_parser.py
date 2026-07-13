"""
Document parser for PDF, DOCX, and TXT files.
Extracts text content from uploaded legal documents.
"""

import io
from typing import Optional


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"


def parse_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            return f"Error parsing text file: {str(e)}"


def parse_document(uploaded_file) -> Optional[str]:
    """
    Parse an uploaded Streamlit file object.
    Supports PDF, DOCX, and TXT formats.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Extracted text content or None on failure
    """
    if uploaded_file is None:
        return None

    file_bytes = uploaded_file.getvalue()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return parse_docx(file_bytes)
    elif filename.endswith(".txt"):
        return parse_txt(file_bytes)
    else:
        return f"Unsupported file format: {filename}"


def get_file_type(filename: str) -> str:
    """Get the file type from filename."""
    filename = filename.lower()
    if filename.endswith(".pdf"):
        return "pdf"
    elif filename.endswith(".docx"):
        return "docx"
    elif filename.endswith(".txt"):
        return "txt"
    return "unknown"
